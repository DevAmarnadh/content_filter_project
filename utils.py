import os
from typing import Tuple, List
import docx
import fitz  # PyMuPDF
from PIL import Image
import io
import re

def get_document_type(file_path: str) -> str:
    """
    Determine the type of document based on file extension.
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.docx':
        return 'docx'
    elif ext == '.pdf':
        return 'pdf'
    elif ext == '.txt':
        return 'txt'
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def extract_docx_content(file_path: str) -> Tuple[List[str], List[Image.Image]]:
    """
    Extract text and images from a .docx file.
    """
    doc = docx.Document(file_path)
    texts = []
    images = []
    
    # Extract text
    for paragraph in doc.paragraphs:
        texts.append(paragraph.text)
    
    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))
            images.append(image)
    
    return texts, images

def extract_pdf_content(file_path: str) -> Tuple[List[str], List[Image.Image]]:
    """
    Extract text and images from a PDF file with improved content detection.
    """
    doc = fitz.open(file_path)
    texts = []
    images = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Extract text with better formatting preservation
        blocks = page.get_text("blocks")
        for block in blocks:
            # Extract text content from block
            text = block[4]
            # Clean the text
            text = re.sub(r'\s+', ' ', text).strip()
            if text:
                texts.append(text)
        
        # Extract images with better quality
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Get image format
            image_format = base_image["ext"].upper()
            if image_format == "JPG":
                image_format = "JPEG"
            
            try:
                # Convert image bytes to PIL Image
                image = Image.open(io.BytesIO(image_bytes))
                # Convert to RGB if necessary
                if image.mode not in ('RGB', 'L'):
                    image = image.convert('RGB')
                images.append(image)
            except Exception as e:
                print(f"Warning: Could not process image {img_index} on page {page_num + 1}: {str(e)}")
                continue
    
    return texts, images

def extract_txt_content(file_path: str) -> Tuple[List[str], List[Image.Image]]:
    """
    Extract text from a .txt file (no images).
    """
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.readlines()
    return text, []

def save_docx(texts: List[str], images: List[Image.Image], output_path: str):
    """
    Save processed content to a new .docx file.
    """
    doc = docx.Document()
    
    # Add text
    for text in texts:
        if text.strip():  # Only add non-empty paragraphs
            doc.add_paragraph(text)
    
    # Add images
    for img in images:
        with io.BytesIO() as bio:
            img.save(bio, format='PNG')
            doc.add_picture(bio)
    
    doc.save(output_path)

def save_pdf(texts: List[str], images: List[Image.Image], output_path: str):
    """
    Save processed content to a new PDF file with better formatting.
    """
    doc = fitz.open()
    current_page = doc.new_page()
    
    # PDF formatting parameters
    margin_x = 50
    margin_y = 50
    y_position = margin_y
    page_height = current_page.rect.height
    page_width = current_page.rect.width
    line_height = 15
    
    # Add text with proper formatting
    for text in texts:
        if not text.strip():
            continue
            
        # Check if we need a new page
        if y_position + line_height > page_height - margin_y:
            current_page = doc.new_page()
            y_position = margin_y
        
        # Add text with proper formatting
        current_page.insert_text(
            point=(margin_x, y_position),
            text=text,
            fontsize=11,
            fontname="helv"
        )
        y_position += line_height * (text.count('\n') + 1.5)
    
    # Add images
    for img in images:
        # Convert PIL image to bytes
        with io.BytesIO() as bio:
            img.save(bio, format='PNG')
            img_bytes = bio.getvalue()
            
            # Check if we need a new page
            if y_position + 300 > page_height - margin_y:
                current_page = doc.new_page()
                y_position = margin_y
            
            # Calculate image dimensions while maintaining aspect ratio
            img_width = min(page_width - 2 * margin_x, 500)
            img_height = min(300, img_width * img.size[1] / img.size[0])
            
            # Define image rectangle
            img_rect = fitz.Rect(
                margin_x,
                y_position,
                margin_x + img_width,
                y_position + img_height
            )
            
            # Insert image
            current_page.insert_image(img_rect, stream=img_bytes)
            y_position += img_height + line_height
    
    doc.save(output_path)
    doc.close()

def save_txt(texts: List[str], output_path: str):
    """
    Save processed text to a new .txt file.
    """
    with open(output_path, 'w', encoding='utf-8') as file:
        for text in texts:
            if text.strip():  # Only write non-empty lines
                file.write(text + '\n') 